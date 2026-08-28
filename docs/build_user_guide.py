#!/usr/bin/env python3
"""Build the Org Explorer user guide."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Twips

NAVY = RGBColor(0x1F, 0x38, 0x64)
ACCENT = RGBColor(0xE3, 0x52, 0x05)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BODY = RGBColor(0x1A, 0x23, 0x32)
MUTED = RGBColor(0x5B, 0x6B, 0x82)
HEADER_FILL = "1F3864"
PLACEHOLDER_FILL = "F3EBE4"
TABLE_WIDTH = 9026

OUT = Path(__file__).with_name("Org-Explorer_User-Guide_v1.2.docx")


def set_run_font(run, *, size=11, bold=False, color=BODY, name="Calibri"):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)


def shade_cell(cell, fill):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_margins(cell, top=60, left=100, bottom=60, right=100):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for edge, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        mar.append(node)
    tcPr.append(mar)


def set_cell_width(cell, width):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:w"), str(width))
    tcW.set(qn("w:type"), "dxa")


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "1F3864")
        borders.append(el)
    tblPr.append(borders)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(TABLE_WIDTH))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)


def prevent_row_split(row):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    cant = OxmlElement("w:cantSplit")
    trPr.append(cant)


def header_row(row):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    trPr.append(tblHeader)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.autofit = False
    set_table_borders(table)
    hdr = table.rows[0]
    header_row(hdr)
    prevent_row_split(hdr)
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        shade_cell(cell, HEADER_FILL)
        set_cell_margins(cell)
        set_cell_width(cell, widths[i])
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text)
        set_run_font(run, size=10, bold=True, color=WHITE)
    for r_i, row_data in enumerate(rows):
        row = table.rows[r_i + 1]
        prevent_row_split(row)
        for c_i, text in enumerate(row_data):
            cell = row.cells[c_i]
            set_cell_margins(cell)
            set_cell_width(cell, widths[c_i])
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(text)
            set_run_font(run, size=10)
    doc.add_paragraph()
    return table


def para(doc, text="", *, after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    if text:
        run = p.add_run(text)
        set_run_font(run, size=11)
    return p


def runs(doc, parts, *, after=8):
    p = para(doc, after=after)
    for part in parts:
        if isinstance(part, str):
            run = p.add_run(part)
            set_run_font(run, size=11)
        else:
            run = p.add_run(part["text"])
            set_run_font(run, size=11, bold=part.get("bold", False), color=part.get("color", BODY))
    return p


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = NAVY
        run.font.name = "Calibri"
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:ascii"), "Calibri")
        rFonts.set(qn("w:hAnsi"), "Calibri")
    if level == 1:
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(8)
    else:
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    return p


def bullet(doc, parts):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.clear()
    for part in parts if isinstance(parts, list) else [{"text": parts}]:
        if isinstance(part, str):
            run = p.add_run(part)
            set_run_font(run, size=11)
        else:
            run = p.add_run(part["text"])
            set_run_font(run, size=11, bold=part.get("bold", False))
    return p


def numbered(doc, parts):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.clear()
    for part in parts if isinstance(parts, list) else [{"text": parts}]:
        if isinstance(part, str):
            run = p.add_run(part)
            set_run_font(run, size=11)
        else:
            run = p.add_run(part["text"])
            set_run_font(run, size=11, bold=part.get("bold", False))
    return p


def placeholder(doc, caption):
    table = doc.add_table(rows=1, cols=1)
    set_table_borders(table)
    cell = table.rows[0].cells[0]
    shade_cell(cell, PLACEHOLDER_FILL)
    set_cell_margins(cell, top=120, left=140, bottom=120, right=140)
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run("Screenshot placeholder")
    set_run_font(run, size=10, bold=True, color=ACCENT)
    p2 = cell.add_paragraph()
    run2 = p2.add_run(f"Insert: {caption}")
    set_run_font(run2, size=10, color=MUTED)
    doc.add_paragraph()


def add_toc(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    fld.set(qn("w:dirty"), "true")
    run._r.append(fld)
    run2 = p.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-2" \\h \\z \\u '
    run2._r.append(instr)
    run3 = p.add_run()
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    run3._r.append(sep)
    run4 = p.add_run("(The contents list updates when you open this file in Word.)")
    set_run_font(run4, size=10, color=MUTED)
    run5 = p.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run5._r.append(end)


def set_header_footer(doc):
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.clear()
    run = hp.add_run("Org Explorer  ·  User guide")
    set_run_font(run, size=9, color=MUTED)
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.clear()
    run = fp.add_run("Page ")
    set_run_font(run, size=9, color=MUTED)
    # PAGE field
    r = fp.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    r._r.append(fld)
    r2 = fp.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    r2._r.append(instr)
    r3 = fp.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    r3._r.append(end)


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    set_header_footer(doc)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].font.color.rgb = BODY

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("Org Explorer")
    set_run_font(run, size=28, bold=True, color=NAVY)

    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(14)
    run = sub.add_run("User guide")
    set_run_font(run, size=18, color=ACCENT)

    add_table(
        doc,
        ["", ""],
        [
            ["Version", "1.2"],
            ["Date", "28 August 2026"],
            ["Author", "David Ahlhaus"],
            ["Status", "Issued"],
        ],
        [2800, 6226],
    )

    para(doc, "This guide is for everyone at Factor1 who needs to look up people, teams, and reporting lines.")

    heading(doc, "Contents")
    add_toc(doc)

    heading(doc, "1. About this guide")
    para(doc, "This guide shows you how to use Org Explorer. Read it at your desk with the site open.")
    para(doc, "Button names in this guide match the words on screen. They appear in bold.")

    heading(doc, "2. What Org Explorer is")
    para(doc, "Org Explorer is Factor1’s live picture of who reports to whom.")
    para(doc, "You use it to find a person, see their manager and team, and move around the firm.")
    para(doc, "The people and their details come from Microsoft 365. When someone updates a detail in Org Explorer, Microsoft 365 stores that change.")

    heading(doc, "3. How you reach it")
    para(doc, "Open Org Explorer in your web browser:")
    runs(doc, [{"text": "https://org-explorer-ruby.vercel.app/", "bold": True}])
    para(doc, "Sign in with your Factor1 Microsoft account. Use the same sign-in you use for email.")

    heading(doc, "4. Before you start")
    numbered(doc, "Use a computer and a current browser, such as Microsoft Edge or Chrome.")
    numbered(doc, "Have your Factor1 Microsoft password ready, or stay signed in to Microsoft already.")
    numbered(doc, "Open the address in Section 3 and choose Sign in.")
    para(doc, "Innovation & Systems turns on Admin for the people who update details. If you need Admin, ask them.")

    heading(doc, "5. Key ideas")
    add_table(
        doc,
        ["Word", "What it means"],
        [
            ["Branch", "The slice of the firm you are looking at: a person in the centre, their manager above, people at the same level beside them, and the people who report to them below."],
            ["Card", "The box for one person. It shows their photo, name, role, company, department, and team."],
            ["Company", "The Factor1 business the person belongs to, such as Factor1 Shepp or Taxopia."],
            ["Department", "The area of work, such as Service Delivery or Firm Management."],
            ["Team", "The named group inside a department, such as Zeus or Avengers."],
            ["Role", "The person’s job title."],
            ["Reports to", "The person this person reports to. Microsoft 365 stores this as the manager."],
            ["Full tree", "A zoomable view of everyone, or of the company–department–team–role structure."],
            ["Admin", "The screens and buttons for people who update company, department, team, role, and Reports to."],
        ],
        [2200, 6826],
    )

    heading(doc, "6. Who can do what")
    add_table(
        doc,
        ["Who", "What they can do"],
        [
            ["Everyone who signs in", "Search, filter, open a branch, and open Full tree."],
            ["People with Admin", "Everything above, plus edit a person’s company, department, team, role, and Reports to. They can also download a list of everyone and apply an updated list."],
        ],
        [2800, 6226],
    )
    para(doc, "Innovation & Systems decides who has Admin. The Admin button appears at the top of the screen when you have it.")

    heading(doc, "7. Getting started")
    numbered(doc, [{"text": "Open "}, {"text": "https://org-explorer-ruby.vercel.app/", "bold": True}, {"text": "."}])
    numbered(doc, [{"text": "Choose "}, {"text": "Sign in", "bold": True}, {"text": "."}])
    numbered(doc, "Pick your Factor1 Microsoft account and finish the Microsoft prompts.")
    numbered(doc, "You land on your own branch. Your card sits in the centre.")
    placeholder(doc, "the Org Explorer screen after sign-in, with your card in the centre, the manager above, and people who report to you below.")
    para(doc, "The top of the screen shows Sign out, My profile, and Full tree. Admin appears only if you have it.")

    heading(doc, "8. Finding people")
    heading(doc, "8.1 Search", 2)
    para(doc, "Type a name in Search by name… at the top. A list of matches appears as you type. Choose a name to open that person’s branch.")
    heading(doc, "8.2 Filter", 2)
    para(doc, "Use the three lists to the right of search:")
    bullet(doc, [{"text": "All companies", "bold": True}, {"text": " — show one company."}])
    bullet(doc, [{"text": "All departments", "bold": True}, {"text": " — show one department. This list follows the company you picked."}])
    bullet(doc, [{"text": "All teams", "bold": True}, {"text": " — show one team. This list follows the company and department you picked."}])
    para(doc, "Set a list back to All… when you want to see everyone again.")
    placeholder(doc, "the top bar with Search by name and the company, department, and team lists.")

    heading(doc, "9. Moving around a branch")
    para(doc, "This is the main way you use Org Explorer.")
    numbered(doc, "Your card is in the centre. The person you report to is above you. People at your level sit beside you. People who report to you sit below you.")
    numbered(doc, [{"text": "To open someone else’s branch, choose their card. On a manager’s card this is labelled "}, {"text": "Go up", "bold": True}, {"text": ". On a peer’s card this is labelled "}, {"text": "View branch", "bold": True}, {"text": ". On a report’s card you see how many people report to them."}])
    numbered(doc, "A path of names appears above the cards after you move. Choose a name in that path to jump back to that person.")
    numbered(doc, [{"text": "Choose "}, {"text": "My profile", "bold": True}, {"text": " at any time to return to your own branch."}])
    placeholder(doc, "a branch with the name path above the cards, after you have clicked through to another person.")

    heading(doc, "10. Other tasks")
    heading(doc, "10.1 See everyone at once", 2)
    numbered(doc, [{"text": "Choose "}, {"text": "Full tree", "bold": True}, {"text": "."}])
    numbered(doc, [{"text": "Choose "}, {"text": "Employees", "bold": True}, {"text": " to see people, or "}, {"text": "Structure", "bold": True}, {"text": " to see company, then department, then team, then role."}])
    numbered(doc, "Drag to move around. Scroll to zoom.")
    numbered(doc, [{"text": "Choose "}, {"text": "Focus me", "bold": True}, {"text": " to jump to your own card."}])
    numbered(doc, [{"text": "Choose "}, {"text": "Print", "bold": True}, {"text": ". The chart prints in landscape, with one branch per page so names stay readable. Each page names the branch. Set company, department, or team first if you only need part of the firm. In the print dialog you can Save as PDF."}])
    numbered(doc, [{"text": "Choose "}, {"text": "Download SVG", "bold": True}, {"text": " to save the whole chart as a file you can open and zoom."}])
    numbered(doc, [{"text": "Choose "}, {"text": "Back to explorer", "bold": True}, {"text": " to return to the branch view."}])
    placeholder(doc, "Full tree in Employees view, with Focus me, Print, Download SVG, and Back to explorer visible.")

    heading(doc, "10.2 Sign out", 2)
    runs(doc, ["Choose ", {"text": "Sign out", "bold": True}, {"text": " when you finish. Sign in again next time."}])

    heading(doc, "10.3 Update a person’s details", 2)
    para(doc, "This section is for people with Admin.")
    para(doc, "You can change Reports to, company, department, team, and role. Microsoft 365 stores the change.")
    numbered(doc, [{"text": "Choose "}, {"text": "Admin", "bold": True}, {"text": ", or choose the pencil on a person’s card."}])
    numbered(doc, [{"text": "In Admin, type a name in "}, {"text": "Search people to edit…", "bold": True}, {"text": " and choose "}, {"text": "Edit", "bold": True}, {"text": "."}])
    numbered(doc, [{"text": "Under "}, {"text": "Reports to", "bold": True}, {"text": ", search for a person and choose their name. Use × if this person sits at the top of the org."}])
    numbered(doc, "Pick a value from each of the company, department, team, and role lists. Use — None — to leave that field blank.")
    numbered(doc, [{"text": "To create a value that is not in a list, choose "}, {"text": "Add new value…", "bold": True}, {"text": ", type the name, then choose "}, {"text": "Add", "bold": True}, {"text": "."}])
    numbered(doc, [{"text": "Choose "}, {"text": "Save changes", "bold": True}, {"text": ". The button shows Saving… while Microsoft 365 stores the change and the chart updates. The panel then shows “Saved to Microsoft 365.”"}])
    numbered(doc, [{"text": "Choose "}, {"text": "Cancel", "bold": True}, {"text": " or the × to close the panel and see the branch."}])
    placeholder(doc, "the Edit details panel with Reports to at the top, the four lists, and Save changes.")
    para(doc, "A company, department, team, or role stays in the lists while at least one person has it. When the last person moves off a value, that value leaves the lists.")
    para(doc, "Org Explorer stops a reporting line that would loop, including a person reporting to themselves.")

    heading(doc, "10.4 Update many people at once", 2)
    para(doc, "This section is for people with Admin.")
    numbered(doc, [{"text": "In Admin, choose "}, {"text": "Download CSV", "bold": True}, {"text": ". A file called org-users.csv saves to your computer."}])
    numbered(doc, "Open the file in Excel.")
    numbered(doc, "Change only the company, department, team, and role columns. Leave the id column as it is, so each row still matches the right person.")
    numbered(doc, "Save the file.")
    numbered(doc, [{"text": "In Admin, choose "}, {"text": "Apply CSV", "bold": True}, {"text": ", pick your file, and confirm."}])
    para(doc, "Wait until the status line at the top reports how many people were updated. If some rows fail, the status line explains why.")

    heading(doc, "11. Points to remember")
    bullet(doc, "Sign in with your Factor1 Microsoft account, not a personal Microsoft account.")
    bullet(doc, [{"text": "My profile", "bold": True}, {"text": " always takes you back to your own branch."}])
    bullet(doc, "Search looks up names. Filters narrow the chart by company, department, and team.")
    bullet(doc, "People with Admin pick company, department, team, and role from lists. They create a new value with Add new value…, then Add.")
    bullet(doc, [{"text": "Reports to", "bold": True}, {"text": " is a person search. Use × if this person sits at the top of the org."}])
    bullet(doc, "Finish Add, or pick an existing value, before Save changes.")
    bullet(doc, [{"text": "Save changes", "bold": True}, {"text": " shows Saving… while Microsoft 365 stores the change and the chart updates."}])
    bullet(doc, [{"text": "Print", "bold": True}, {"text": " on Full tree puts each branch on its own landscape page. "}, {"text": "Download SVG", "bold": True}, {"text": " saves the whole chart as a file you can zoom."}])
    bullet(doc, "When you apply a list, leave each person’s id unchanged.")
    bullet(doc, "A change you save in Org Explorer is the change Microsoft 365 keeps.")

    heading(doc, "12. If something goes wrong")
    add_table(
        doc,
        ["What you see", "What to do"],
        [
            ["Microsoft asks you to sign in again, or sign-in fails.", "Use your Factor1 account. Try a private window, or a different browser. If it still fails, contact Innovation & Systems."],
            ["The screen stays on Loading…", "Refresh the page. If it continues, contact Innovation & Systems."],
            ["Admin is missing from the top of the screen.", "Ask Innovation & Systems if you need it. The button appears only for people with Admin."],
            ["Save changes reports a problem.", "Read the message on the panel. Try again. If it repeats, send the message to Innovation & Systems."],
            ["The panel reports Reports to could not be updated.", "Close the panel, open the person’s branch, and check who sits above them. If the old reporting line is still there, send the on-screen message to Innovation & Systems."],
            ["Apply CSV reports that some people failed.", "Read the status line. Send that text to Innovation & Systems, with the file you applied."],
            ["A person’s team did not update.", "Tell Innovation & Systems the person’s name. Some mailboxes need a change in Exchange instead."],
            ["A person is missing from the chart.", "Confirm they have an active Factor1 Microsoft account. If they do, tell Innovation & Systems."],
        ],
        [3200, 5826],
    )

    heading(doc, "13. Glossary")
    add_table(
        doc,
        ["Term", "Meaning"],
        [
            ["Admin", "The area of Org Explorer used to update company, department, team, role, and Reports to."],
            ["Branch", "The view centred on one person and the people directly around them."],
            ["Card", "The on-screen box for one person."],
            ["Company", "The Factor1 business a person belongs to."],
            ["Department", "The area of work a person belongs to."],
            ["Full tree", "The zoomable view of everyone, or of the structure of the firm."],
            ["Microsoft 365", "The firm’s directory of people. Org Explorer reads from it and, for Admin, writes back to it."],
            ["Reports to", "The person this person reports to. Microsoft 365 stores this as the manager."],
            ["Role", "A person’s job title."],
            ["Team", "The named group inside a department."],
        ],
        [2200, 6826],
    )

    heading(doc, "14. Contacts")
    add_table(
        doc,
        ["Question", "Who to ask"],
        [
            ["How do I use Org Explorer?", "David Ahlhaus, Innovation & Systems"],
            ["I need Admin.", "Innovation & Systems"],
            ["A save or list update failed.", "Innovation & Systems — include the on-screen message"],
            ["Someone is missing or in the wrong place.", "Innovation & Systems"],
        ],
        [3600, 5426],
    )

    heading(doc, "15. Version history")
    add_table(
        doc,
        ["Version", "Date", "Author", "Changes"],
        [
            ["1.0", "27 August 2026", "David Ahlhaus", "First issue. Covers sign-in, search, filters, branches, Full tree, and Admin updates."],
            ["1.1", "28 August 2026", "David Ahlhaus", "Added Reports to on Edit details, Saving… on Save changes, and Print on Full tree."],
            ["1.2", "28 August 2026", "David Ahlhaus", "Print on Full tree now uses one landscape page per branch. Download SVG saves the whole chart as a file you can zoom."],
        ],
        [1400, 1800, 2200, 3626],
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
